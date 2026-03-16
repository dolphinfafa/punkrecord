"""
Project API endpoints
"""
from typing import Optional, Dict, Any, List, Union, Tuple
from uuid import UUID, uuid4
from datetime import datetime
import io
import urllib.parse
from pathlib import Path

from fastapi import APIRouter, Depends, Query, Body, UploadFile, File
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from sqlmodel import Session, select
from app.core.database import get_session
from app.core.auth import get_current_user, require_permission
from app.core.exceptions import NotFoundException, ForbiddenException, ValidationException
from app.core.response import success_response
from app.models.iam import User, OurEntity, JobTitle
from app.models.project import Project, ProjectStage, ProjectMember, ProjectType, ProjectStatus, StageStatus
from app.models.todo import TodoItem, TodoSourceType, TodoPriority
from app.models.contract import Counterparty
from app.services.docx_generator import generate_acceptance_report
from app.core.storage import save_file, get_file, delete_file, get_download_url
from app.schemas.project import (
    ProjectCreate, ProjectUpdate, ProjectResponse,
    ProjectStageResponse, StageStatusUpdate,
    ProjectMemberCreate, ProjectMemberBatchCreate, ProjectMemberResponse, ProjectTaskResponse,
    ProjectTaskAssignRequest, ProjectTaskPlanUpdateRequest, ProjectTaskBatchAssignRequest
)

router = APIRouter(prefix="/project", tags=["Project"])
PROJECT_ATTACHMENT_DIR = Path(__file__).resolve().parents[2] / "uploads" / "project-attachments"
MAX_PROJECT_ATTACHMENT_SIZE = 20 * 1024 * 1024  # 20 MB per file
STAGE_ATTACHMENT_DIR = Path(__file__).resolve().parents[2] / "uploads" / "project-stage-attachments"
MAX_STAGE_ATTACHMENT_SIZE = 20 * 1024 * 1024  # 20 MB per file
DEV_TYPE_LABELS = {
    "dev_backend": "\u540e\u7aef",
    "dev_frontend": "\u524d\u7aef",
    "dev_ui": "UI",
    "dev_product": "\u4ea7\u54c1",
}
VALID_DEV_TYPES = set(DEV_TYPE_LABELS.keys()) | {"other"}


# Standard stages for B2B projects
B2B_STAGES = [
    ("requirement_alignment", "需求对齐"),
    ("quotation", "报价"),
    ("contract_signed", "合同签署"),
    ("prototype_confirmed", "原型确认"),
    ("design", "设计"),
    ("development", "开发"),
    ("testing", "测试"),
    ("delivery", "交付")
]

# Standard stages for B2C projects
B2C_STAGES = [
    ("requirement_analysis", "需求分析"),
    ("project_initiation", "项目立项"),
    ("prototype_confirmed", "原型确认"),
    ("design", "设计"),
    ("development", "开发"),
    ("testing", "测试"),
    ("launch", "上线"),
    ("operation", "运营"),
    ("iteration", "迭代")
]


def enrich_project_response(session: Session, project: Project) -> ProjectResponse:
    resp = ProjectResponse.model_validate(project)
    if project.pm_user_id:
        pm = session.get(User, project.pm_user_id)
        if pm:
            resp.pm_name = pm.display_name or pm.email
    return resp


def _extract_feature_list_for_project(project_id: UUID, session: Session) -> Tuple[List[dict], Optional[ProjectStage]]:
    """Get parsed feature list and its source stage for a project."""
    import json as _json

    stages = session.exec(
        select(ProjectStage)
        .where(ProjectStage.project_id == project_id)
        .order_by(ProjectStage.sequence_no)
    ).all()

    # Priority order:
    # - B2B feature list usually comes from requirement_alignment
    # - B2C feature list usually comes from project_initiation
    # - quotation may contain a different row shape, so keep it lower priority
    priority_codes = ["requirement_alignment", "project_initiation", "prototype_confirmed", "quotation"]
    stage_map = {s.stage_code: s for s in stages}

    def _is_feature_list(data):
        if not isinstance(data, list) or len(data) == 0:
            return False
        first = data[0]
        if not isinstance(first, dict):
            return False
        feature_keys = {"l1_feature", "l2_feature", "module", "dev_backend", "dev_frontend"}
        return bool(feature_keys & set(first.keys()))

    for code in priority_codes:
        stage = stage_map.get(code)
        if not stage or not stage.feature_list:
            continue
        try:
            candidate = _json.loads(stage.feature_list)
            if _is_feature_list(candidate):
                return candidate, stage
        except Exception:
            continue

    for stage in stages:
        if not stage.feature_list:
            continue
        try:
            candidate = _json.loads(stage.feature_list)
            if _is_feature_list(candidate):
                return candidate, stage
        except Exception:
            continue

    return [], None


def _feature_row_dev_types(row: dict) -> List[str]:
    types = []
    for key in ["dev_backend", "dev_frontend", "dev_ui", "dev_product"]:
        value = str(row.get(key, "")).strip()
        if value and value != "-" and value != "0":
            types.append(key)
    # Keep behavior consistent with frontend generation: rows with no dev split default to backend.
    return types if types else ["dev_backend"]


def _build_feature_task_title(row: dict, dev_type: str) -> str:
    label = DEV_TYPE_LABELS.get(dev_type, dev_type)
    name = " - ".join([x for x in [row.get("l1_feature"), row.get("l2_feature")] if x])
    return f"[{label}] {name or '开发任务'}"


def _is_project_member(session: Session, project: Project, user_id: UUID) -> bool:
    if user_id in {project.pm_user_id, project.owner_user_id}:
        return True
    member = session.exec(
        select(ProjectMember)
        .where(ProjectMember.project_id == project.id)
        .where(ProjectMember.user_id == user_id)
    ).first()
    return bool(member)


def _is_bug_todo(todo: TodoItem) -> bool:
    tags = list(todo.tags or [])
    link = dict(todo.link or {})
    return "bug" in tags or link.get("type") == "bug"


def _normalize_ref_id(value: Any) -> str:
    return str(value or "").strip().lower()


def _collect_related_todos_for_project_todo(session: Session, todo: TodoItem) -> List[TodoItem]:
    """Find mirrored/linked todos that should be deleted together."""
    all_todos = session.exec(select(TodoItem)).all()
    target_id = _normalize_ref_id(todo.id)
    link = dict(todo.link or {})
    direct_ref = _normalize_ref_id(link.get("tracking_todo_id") or link.get("project_todo_id"))

    related: List[TodoItem] = []
    seen_ids = {target_id}
    for item in all_todos:
        item_id = _normalize_ref_id(item.id)
        if item_id in seen_ids:
            continue
        item_link = dict(item.link or {})
        item_ref = _normalize_ref_id(item_link.get("tracking_todo_id") or item_link.get("project_todo_id"))
        if item_ref and item_ref == target_id:
            related.append(item)
            seen_ids.add(item_id)
            continue
        if direct_ref and item_id == direct_ref:
            related.append(item)
            seen_ids.add(item_id)
    return related


def _delete_project_todo_with_related(session: Session, todo: TodoItem) -> int:
    count = 0
    for related in _collect_related_todos_for_project_todo(session, todo):
        session.delete(related)
        count += 1
    session.delete(todo)
    count += 1
    return count


@router.post("/projects", response_model=dict)
async def create_project(
    data: ProjectCreate,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Create project with automatic stage generation"""
    # Fetch default our_entity if not provided
    our_entity_id = data.our_entity_id
    if not our_entity_id:
        default_entity = session.exec(select(OurEntity)).first()
        if default_entity:
            our_entity_id = default_entity.id
        else:
            raise NotFoundException("未找到可用的我方主体")
            
    project = Project(
        our_entity_id=our_entity_id,
        project_no=data.project_no,
        name=data.name,
        project_type=ProjectType(data.project_type),
        status=ProjectStatus.DRAFT,
        owner_user_id=current_user.id,
        pm_user_id=data.pm_user_id,
        customer_id=data.customer_id,
        contract_id=data.contract_id,
        start_at=data.start_at,
        due_at=data.due_at,
        current_stage_code="",
        progress=0.0,
        description=data.description
    )
    
    session.add(project)
    session.commit()
    session.refresh(project)
    
    # Generate stages based on project type
    stages = B2B_STAGES if project.project_type == ProjectType.B2B else B2C_STAGES
    
    for idx, (code, name) in enumerate(stages, 1):
        stage = ProjectStage(
            project_id=project.id,
            stage_code=code,
            stage_name=name,
            sequence_no=idx,
            status=StageStatus.NOT_STARTED
        )
        session.add(stage)
    
    # Set current stage to first stage
    project.current_stage_code = stages[0][0]
    session.add(project)
    session.commit()
    session.refresh(project)
    return success_response(enrich_project_response(session, project))

@router.get("/projects", response_model=dict)
async def list_projects(
    status: Optional[str] = Query(None),
    project_type: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """List projects"""
    query = select(Project)
    
    if status:
        query = query.where(Project.status == status)
    if project_type:
        query = query.where(Project.project_type == project_type)
    
    query = query.order_by(Project.created_at.desc())
    
    offset = (page - 1) * page_size
    projects = session.exec(query.offset(offset).limit(page_size)).all()
    
    count_query = select(Project)
    if status:
        count_query = count_query.where(Project.status == status)
    if project_type:
        count_query = count_query.where(Project.project_type == project_type)
    total = len(session.exec(count_query).all())
    
    return success_response({
        "items": [enrich_project_response(session, p) for p in projects],
        "total": total,
        "page": page,
        "page_size": page_size,
        "pages": (total + page_size - 1) // page_size
    })


@router.get("/projects/{project_id}", response_model=dict)
async def get_project(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """Get project by ID"""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
    
    return success_response(enrich_project_response(session, project))


@router.get("/projects/{project_id}/attachments", response_model=dict)
async def list_project_attachments(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """List project-level attachments."""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    attachments = list(project.attachments or [])
    return success_response(attachments)


@router.post("/projects/{project_id}/attachments", response_model=dict)
async def upload_project_attachment(
    project_id: UUID,
    file: UploadFile = File(...),
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Upload an attachment for the project."""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    if not file.filename:
        raise ValidationException("附件文件名不能为空")

    file_bytes = await file.read()
    if not file_bytes:
        raise ValidationException("附件内容不能为空")
    if len(file_bytes) > MAX_PROJECT_ATTACHMENT_SIZE:
        raise ValidationException("附件大小不能超过 20MB")

    attachment_id = uuid4().hex
    safe_suffix = Path(file.filename).suffix[:20]
    stored_name = f"{project_id}_{attachment_id}{safe_suffix}"
    save_file("project-attachments", stored_name, file_bytes, file.content_type or "application/octet-stream")

    attachment = {
        "id": attachment_id,
        "file_name": file.filename,
        "stored_name": stored_name,
        "content_type": file.content_type or "application/octet-stream",
        "size": len(file_bytes),
        "uploaded_at": datetime.utcnow().isoformat()
    }
    attachments = list(project.attachments or [])
    attachments.append(attachment)
    project.attachments = attachments
    project.updated_at = datetime.utcnow()
    session.add(project)
    session.commit()

    return success_response(attachment)


@router.get("/projects/{project_id}/attachments/{attachment_id}/download")
async def download_project_attachment(
    project_id: UUID,
    attachment_id: str,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """Download a project-level attachment."""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    attachment = next((a for a in (project.attachments or []) if a.get("id") == attachment_id), None)
    if not attachment:
        raise NotFoundException("未找到附件")

    stored_name = attachment.get("stored_name")
    file_name = attachment.get("file_name") or "attachment"
    if not stored_name:
        raise NotFoundException("附件元数据异常")

    # TOS: redirect to pre-signed URL; local: return file directly
    url = get_download_url("project-attachments", stored_name)
    if url:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url)

    try:
        _, local_path = get_file("project-attachments", stored_name)
    except FileNotFoundError:
        raise NotFoundException("附件文件不存在")

    return FileResponse(
        path=local_path,
        media_type=attachment.get("content_type") or "application/octet-stream",
        filename=file_name
    )


@router.delete("/projects/{project_id}/attachments/{attachment_id}", response_model=dict)
async def delete_project_attachment(
    project_id: UUID,
    attachment_id: str,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Delete a project-level attachment."""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    attachments = list(project.attachments or [])
    idx = next((i for i, item in enumerate(attachments) if item.get("id") == attachment_id), None)
    if idx is None:
        raise NotFoundException("未找到附件")

    removed = attachments.pop(idx)
    stored_name = removed.get("stored_name")
    if stored_name:
        try:
            delete_file("project-attachments", stored_name)
        except Exception:
            pass  # best-effort cleanup

    project.attachments = attachments
    project.updated_at = datetime.utcnow()
    session.add(project)
    session.commit()

    return success_response({"message": "附件已删除"})


@router.patch("/projects/{project_id}", response_model=dict)
async def update_project(
    project_id: UUID,
    data: ProjectUpdate,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Update project"""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
    
    if data.name is not None:
        project.name = data.name
    if data.status is not None:
        project.status = ProjectStatus(data.status)
    if data.pm_user_id is not None:
        project.pm_user_id = data.pm_user_id
    if data.description is not None:
        project.description = data.description
    if data.start_at is not None:
        project.start_at = data.start_at
    if data.due_at is not None:
        project.due_at = data.due_at
    if data.our_entity_id is not None:
        project.our_entity_id = data.our_entity_id
    if data.customer_id is not None:
        project.customer_id = data.customer_id
    
    project.updated_at = datetime.utcnow()
    session.add(project)
    session.commit()
    session.refresh(project)
    
    return success_response(enrich_project_response(session, project))


@router.get("/projects/{project_id}/stages", response_model=dict)
async def get_project_stages(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """Get project stages"""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
    
    stages = session.exec(
        select(ProjectStage)
        .where(ProjectStage.project_id == project_id)
        .order_by(ProjectStage.sequence_no)
    ).all()

    for stage in stages:
        if stage.attachments is None:
            stage.attachments = []

    return success_response([ProjectStageResponse.model_validate(s) for s in stages])


@router.patch("/projects/{project_id}/stages/{stage_id}", response_model=dict)
async def update_stage_status(
    project_id: UUID,
    stage_id: UUID,
    data: StageStatusUpdate,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Update project stage (status, deliverables, feature_list)"""
    stage = session.get(ProjectStage, stage_id)
    if not stage or stage.project_id != project_id:
        raise NotFoundException("未找到阶段")

    if data.status is not None:
        stage.status = StageStatus(data.status)
        if data.blocked_reason:
            stage.blocked_reason = data.blocked_reason
        if data.skip_reason:
            stage.skip_reason = data.skip_reason
        if stage.status == StageStatus.IN_PROGRESS and not stage.actual_start_at:
            stage.actual_start_at = datetime.utcnow().date()
        elif stage.status == StageStatus.DONE and not stage.actual_end_at:
            stage.actual_end_at = datetime.utcnow().date()

    if data.deliverables is not None:
        stage.deliverables = data.deliverables
    if data.feature_list is not None:
        stage.feature_list = data.feature_list
    if data.planned_start_at is not None:
        stage.planned_start_at = data.planned_start_at
    if data.planned_end_at is not None:
        stage.planned_end_at = data.planned_end_at

    stage.updated_at = datetime.utcnow()
    session.add(stage)
    session.commit()
    session.refresh(stage)

    if stage.attachments is None:
        stage.attachments = []
    return success_response(ProjectStageResponse.model_validate(stage))


@router.post("/projects/{project_id}/stages/{stage_id}/attachments", response_model=dict)
async def upload_stage_attachment(
    project_id: UUID,
    stage_id: UUID,
    file: UploadFile = File(...),
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Upload a stage attachment and append metadata into stage.attachments."""
    stage = session.get(ProjectStage, stage_id)
    if not stage or stage.project_id != project_id:
        raise NotFoundException("未找到阶段")

    if not file.filename:
        raise ValidationException("附件文件名不能为空")

    file_bytes = await file.read()
    if not file_bytes:
        raise ValidationException("附件内容不能为空")
    if len(file_bytes) > MAX_STAGE_ATTACHMENT_SIZE:
        raise ValidationException("附件大小不能超过 20MB")

    safe_suffix = Path(file.filename).suffix[:20]
    attachment_id = uuid4().hex
    stored_name = f"{project_id}_{stage_id}_{attachment_id}{safe_suffix}"
    save_file("project-stage-attachments", stored_name, file_bytes, file.content_type or "application/octet-stream")

    attachment = {
        "id": attachment_id,
        "file_name": file.filename,
        "stored_name": stored_name,
        "content_type": file.content_type or "application/octet-stream",
        "size": len(file_bytes),
        "uploaded_at": datetime.utcnow().isoformat()
    }

    attachments = list(stage.attachments or [])
    attachments.append(attachment)
    stage.attachments = attachments
    stage.updated_at = datetime.utcnow()

    session.add(stage)
    session.commit()
    session.refresh(stage)

    return success_response(attachment)


@router.get("/projects/{project_id}/stages/{stage_id}/attachments/{attachment_id}/download")
async def download_stage_attachment(
    project_id: UUID,
    stage_id: UUID,
    attachment_id: str,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """Download a specific stage attachment."""
    stage = session.get(ProjectStage, stage_id)
    if not stage or stage.project_id != project_id:
        raise NotFoundException("未找到阶段")

    attachment = next((a for a in (stage.attachments or []) if a.get("id") == attachment_id), None)
    if not attachment:
        raise NotFoundException("未找到附件")

    stored_name = attachment.get("stored_name")
    file_name = attachment.get("file_name") or "attachment"
    if not stored_name:
        raise NotFoundException("附件元数据异常")

    url = get_download_url("project-stage-attachments", stored_name)
    if url:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url)

    try:
        _, local_path = get_file("project-stage-attachments", stored_name)
    except FileNotFoundError:
        raise NotFoundException("附件文件不存在")

    return FileResponse(
        path=local_path,
        media_type=attachment.get("content_type") or "application/octet-stream",
        filename=file_name
    )


@router.delete("/projects/{project_id}/stages/{stage_id}/attachments/{attachment_id}", response_model=dict)
async def delete_stage_attachment(
    project_id: UUID,
    stage_id: UUID,
    attachment_id: str,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Delete stage attachment metadata and file."""
    stage = session.get(ProjectStage, stage_id)
    if not stage or stage.project_id != project_id:
        raise NotFoundException("未找到阶段")

    attachments = list(stage.attachments or [])
    idx = next((i for i, item in enumerate(attachments) if item.get("id") == attachment_id), None)
    if idx is None:
        raise NotFoundException("未找到附件")

    removed = attachments.pop(idx)
    stored_name = removed.get("stored_name")
    if stored_name:
        try:
            delete_file("project-stage-attachments", stored_name)
        except Exception:
            pass

    stage.attachments = attachments
    stage.updated_at = datetime.utcnow()
    session.add(stage)
    session.commit()

    return success_response({"message": "附件已删除"})


@router.delete("/projects/{project_id}", response_model=dict)
async def delete_project(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Delete project"""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
    
    # Check if there are related todos or stages?
    # For now, cascading delete is handled by database or we just delete
    # But usually we should check. Pydantic/SQLAlchemy might handle cascade if configured
    # For simplicity, we just delete the project and rely on DB cascade
    
    session.delete(project)
    session.commit()
    
    return success_response({"message": "项目已删除"})


# --- Project Members ---

@router.post("/projects/{project_id}/members", response_model=dict)
async def add_project_member(
    project_id: UUID,
    data: Union[ProjectMemberCreate, ProjectMemberBatchCreate],
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Add member(s) to project"""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    # Backward compatible:
    # - {"user_id": "..."} for single add
    # - {"user_ids": ["...", "..."]} for batch add
    if isinstance(data, ProjectMemberBatchCreate):
        target_user_ids = data.user_ids
        role_in_project = data.role_in_project
    else:
        target_user_ids = [data.user_id]
        role_in_project = data.role_in_project

    members: List[ProjectMemberResponse] = []

    def _resolve_role_name(user_obj: User, incoming_role: Optional[str]) -> str:
        if incoming_role and incoming_role.strip():
            return incoming_role.strip()
        if user_obj and user_obj.job_title_id:
            jt = session.get(JobTitle, user_obj.job_title_id)
            if jt and jt.name:
                return jt.name
        return "项目成员"
    seen = set()
    for user_id in target_user_ids:
        if user_id in seen:
            continue
        seen.add(user_id)

        user = session.get(User, user_id)
        if not user:
            raise NotFoundException("未找到用户")

        existing = session.exec(
            select(ProjectMember)
            .where(ProjectMember.project_id == project_id)
            .where(ProjectMember.user_id == user_id)
        ).first()

        resolved_role = _resolve_role_name(user, role_in_project)

        if existing:
            # Keep existing role when already set; otherwise backfill from org job title.
            if not existing.role_in_project:
                existing.role_in_project = resolved_role
                existing.updated_at = datetime.utcnow()
                session.add(existing)
                session.commit()
                session.refresh(existing)
            resp = ProjectMemberResponse.model_validate(existing)
            resp.user_name = user.display_name
            resp.user_email = user.email
            members.append(resp)
            continue

        member = ProjectMember(
            project_id=project_id,
            user_id=user_id,
            role_in_project=resolved_role
        )

        session.add(member)
        session.commit()
        session.refresh(member)

        resp = ProjectMemberResponse.model_validate(member)
        resp.user_name = user.display_name
        resp.user_email = user.email
        members.append(resp)

    if len(members) == 1 and isinstance(data, ProjectMemberCreate):
        return success_response(members[0])
    return success_response(members)


@router.delete("/projects/{project_id}/members/{user_id}", response_model=dict)
async def remove_project_member(
    project_id: UUID,
    user_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Remove member from project"""
    member = session.exec(
        select(ProjectMember)
        .where(ProjectMember.project_id == project_id)
        .where(ProjectMember.user_id == user_id)
    ).first()
    
    if member:
        session.delete(member)
        session.commit()
        
    return success_response({"message": "成员已移除"})


@router.get("/projects/{project_id}/members", response_model=dict)
async def list_project_members(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """List project members"""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
        
    members = session.exec(
        select(ProjectMember)
        .where(ProjectMember.project_id == project_id)
    ).all()
    
    # Enhance with user info
    result = []
    for m in members:
        user = session.get(User, m.user_id)
        resp = ProjectMemberResponse.model_validate(m)
        if user:
            resp.user_name = user.display_name
            resp.user_email = user.email
            if not resp.role_in_project:
                jt_name = None
                if user.job_title_id:
                    jt = session.get(JobTitle, user.job_title_id)
                    jt_name = jt.name if jt else None
                resp.role_in_project = jt_name or "项目成员"
        result.append(resp)
        
    return success_response(result)


# --- Project Tasks ---

@router.get("/projects/{project_id}/todos", response_model=dict)
async def list_project_todos(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """List project todos"""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
        
    # Find todos linked to this project
    # source_type = PROJECT_TASK and source_id = project_id (as string)
    
    todos = session.exec(
        select(TodoItem)
        .where(TodoItem.source_type == TodoSourceType.PROJECT_TASK)
        .where(TodoItem.source_id == str(project_id))
        .order_by(TodoItem.created_at.desc())
    ).all()
    
    result = []
    for t in todos:
        assignee = session.get(User, t.assignee_user_id)
        creator = session.get(User, t.creator_user_id)
        resp = ProjectTaskResponse.model_validate(t)
        if assignee:
            resp.assignee_name = assignee.display_name
        if creator:
            resp.creator_name = creator.display_name
        result.append(resp)
        
    return success_response(result)


@router.post("/projects/{project_id}/todos/{todo_id}/assign", response_model=dict)
async def assign_project_todo(
    project_id: UUID,
    todo_id: UUID,
    data: ProjectTaskAssignRequest,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Assign/reassign a project task to a member. Only PM can assign."""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    if current_user.id not in {project.pm_user_id, project.owner_user_id}:
        raise ForbiddenException("仅项目经理可分配任务")

    todo = session.get(TodoItem, todo_id)
    if not todo:
        raise NotFoundException("未找到任务")
    if todo.source_type != TodoSourceType.PROJECT_TASK or todo.source_id != str(project_id):
        raise NotFoundException("任务不属于当前项目")

    assignee = session.get(User, data.assignee_user_id)
    if not assignee:
        raise NotFoundException("未找到用户")

    member = session.exec(
        select(ProjectMember)
        .where(ProjectMember.project_id == project_id)
        .where(ProjectMember.user_id == data.assignee_user_id)
    ).first()
    if not member and data.assignee_user_id != project.pm_user_id:
        raise ForbiddenException("只能指派给项目成员或项目经理")

    todo.assignee_user_id = data.assignee_user_id
    todo.updated_at = datetime.utcnow()
    session.add(todo)
    session.commit()
    session.refresh(todo)

    resp = ProjectTaskResponse.model_validate(todo)
    resp.assignee_name = assignee.display_name
    creator = session.get(User, todo.creator_user_id)
    if creator:
        resp.creator_name = creator.display_name
    return success_response(resp)


@router.post("/projects/{project_id}/todos/{todo_id}/plan", response_model=dict)
async def update_project_todo_plan(
    project_id: UUID,
    todo_id: UUID,
    data: ProjectTaskPlanUpdateRequest,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Update project task planning fields.

    PM/owner can update all fields.
    Any project member can modify bug assignee only.
    """
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    todo = session.get(TodoItem, todo_id)
    if not todo:
        raise NotFoundException("未找到任务")
    if todo.source_type != TodoSourceType.PROJECT_TASK or todo.source_id != str(project_id):
        raise NotFoundException("任务不属于当前项目")

    is_pm_or_owner = current_user.id in {project.pm_user_id, project.owner_user_id}
    if not is_pm_or_owner:
        if not _is_bug_todo(todo):
            raise ForbiddenException("仅项目经理可修改开发计划")
        if not _is_project_member(session, project, current_user.id):
            raise ForbiddenException("仅项目组成员可修改 Bug 开发人员")
        if (
            data.assignee_user_id is None
            or data.due_at is not None
            or data.priority is not None
            or data.title is not None
            or data.description is not None
            or data.dev_type is not None
        ):
            raise ForbiddenException("项目组成员仅可修改 Bug 的开发人员")

    if (
        data.assignee_user_id is None
        and data.due_at is None
        and data.priority is None
        and data.title is None
        and data.description is None
        and data.dev_type is None
    ):
        raise ValidationException("至少需要修改一个字段")

    assignee = None
    if data.assignee_user_id is not None:
        assignee = session.get(User, data.assignee_user_id)
        if not assignee:
            raise NotFoundException("未找到用户")

        member = session.exec(
            select(ProjectMember)
            .where(ProjectMember.project_id == project_id)
            .where(ProjectMember.user_id == data.assignee_user_id)
        ).first()
        if not member and data.assignee_user_id != project.pm_user_id:
            raise ForbiddenException("只能指派给项目成员或项目经理")
        todo.assignee_user_id = data.assignee_user_id

    if data.due_at is not None:
        todo.due_at = data.due_at
    if data.priority is not None:
        priority_map = {"p0": TodoPriority.P0, "p1": TodoPriority.P1, "p2": TodoPriority.P2, "p3": TodoPriority.P3}
        priority = priority_map.get(str(data.priority).lower())
        if not priority:
            raise ValidationException("优先级仅支持 p0/p1/p2/p3")
        todo.priority = priority
    if data.title is not None:
        title = data.title.strip()
        if not title:
            raise ValidationException("Task title cannot be empty")
        todo.title = title
    if data.description is not None:
        todo.description = data.description
    if data.dev_type is not None:
        dev_type = data.dev_type.strip()
        if dev_type and dev_type not in VALID_DEV_TYPES:
            raise ValidationException("Invalid dev_type")
        todo.tags = [dev_type] if dev_type else []
        link = dict(todo.link or {})
        link["dev_type"] = dev_type or ""
        todo.link = link

    todo.updated_at = datetime.utcnow()
    session.add(todo)
    session.commit()
    session.refresh(todo)

    resp = ProjectTaskResponse.model_validate(todo)
    final_assignee = assignee or session.get(User, todo.assignee_user_id)
    if final_assignee:
        resp.assignee_name = final_assignee.display_name
    creator = session.get(User, todo.creator_user_id)
    if creator:
        resp.creator_name = creator.display_name
    return success_response(resp)


@router.post("/projects/{project_id}/todos/batch-assign", response_model=dict)
async def batch_assign_project_todos(
    project_id: UUID,
    data: ProjectTaskBatchAssignRequest,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Batch assign project todos to one assignee."""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
    if current_user.id not in {project.pm_user_id, project.owner_user_id}:
        raise ForbiddenException("仅项目经理或项目负责人可批量分配任务")
    if not data.todo_ids:
        raise ValidationException("请选择至少一个任务")

    assignee = session.get(User, data.assignee_user_id)
    if not assignee:
        raise NotFoundException("未找到用户")
    if not _is_project_member(session, project, data.assignee_user_id):
        raise ForbiddenException("只能指派给项目组成员")

    todos = session.exec(
        select(TodoItem)
        .where(TodoItem.id.in_(data.todo_ids))
        .where(TodoItem.source_type == TodoSourceType.PROJECT_TASK)
        .where(TodoItem.source_id == str(project_id))
    ).all()
    if not todos:
        raise ValidationException("未找到可更新的任务")

    now = datetime.utcnow()
    for todo in todos:
        todo.assignee_user_id = data.assignee_user_id
        todo.updated_at = now
        session.add(todo)
    session.commit()
    return success_response({"updated": len(todos), "assignee_user_id": str(data.assignee_user_id)})


@router.delete("/projects/{project_id}/todos/{todo_id}", response_model=dict)
async def delete_project_todo(
    project_id: UUID,
    todo_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Delete a project task/bug todo.

    Allowed: project PM, project owner, or todo creator.
    """
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    todo = session.get(TodoItem, todo_id)
    if not todo:
        raise NotFoundException("未找到任务")
    if todo.source_type != TodoSourceType.PROJECT_TASK or todo.source_id != str(project_id):
        raise NotFoundException("任务不属于当前项目")

    allowed_users = {project.pm_user_id, project.owner_user_id, todo.creator_user_id}
    if current_user.id not in allowed_users:
        raise ForbiddenException("仅项目经理/项目负责人或创建人可删除该任务")

    deleted_count = _delete_project_todo_with_related(session, todo)
    session.commit()
    return success_response({"message": "任务已删除", "deleted": deleted_count})


@router.delete("/projects/{project_id}/todos", response_model=dict)
async def clear_project_todos(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Clear all development tasks in a project and linked mirrored todos."""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
    if current_user.id not in {project.pm_user_id, project.owner_user_id}:
        raise ForbiddenException("仅项目经理或项目负责人可清空任务")

    todos = session.exec(
        select(TodoItem)
        .where(TodoItem.source_type == TodoSourceType.PROJECT_TASK)
        .where(TodoItem.source_id == str(project_id))
    ).all()

    deleted_total = 0
    for todo in todos:
        if session.get(TodoItem, todo.id) is None:
            continue
        deleted_total += _delete_project_todo_with_related(session, todo)
    session.commit()
    return success_response({"deleted": deleted_total})


@router.post("/projects/{project_id}/sync-dev-tasks", response_model=dict)
async def sync_dev_tasks_from_feature_list(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Strictly sync project tasks from feature list so task list matches feature list."""
    from app.models.todo import TodoItem, TodoSourceType, TodoActionType, TodoPriority, TodoStatus

    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("Project not found")
    if current_user.id not in {project.pm_user_id, project.owner_user_id}:
        raise ForbiddenException("Only PM or owner can sync development tasks")

    feature_list, source_stage = _extract_feature_list_for_project(project_id, session)
    if not feature_list:
        return success_response({
            "created": 0,
            "deleted": 0,
            "feature_total": 0,
            "source_stage": None,
            "message": "No feature list found",
        })

    expected_rows: List[Dict[str, Any]] = []
    for idx, row in enumerate(feature_list):
        if not isinstance(row, dict):
            continue
        for dev_type in _feature_row_dev_types(row):
            expected_rows.append({
                "feature_key": f"{idx}|{dev_type}",
                "dev_type": dev_type,
                "title": _build_feature_task_title(row, dev_type),
                "description": row.get("description") or "",
                "l1_feature": row.get("l1_feature"),
                "l2_feature": row.get("l2_feature"),
                "module": row.get("module"),
            })

    existing_todos = session.exec(
        select(TodoItem)
        .where(TodoItem.source_type == TodoSourceType.PROJECT_TASK)
        .where(TodoItem.source_id == str(project_id))
    ).all()

    # Preserve assignee/priority/due for rows that can be matched by feature_key.
    preserved: Dict[str, Dict[str, Any]] = {}
    for todo in existing_todos:
        link = todo.link or {}
        feature_key = link.get("feature_key")
        if not feature_key:
            continue
        preserved[feature_key] = {
            "assignee_user_id": todo.assignee_user_id,
            "priority": todo.priority,
            "due_at": todo.due_at,
        }

    deleted = 0
    for todo in existing_todos:
        if session.get(TodoItem, todo.id) is None:
            continue
        deleted += _delete_project_todo_with_related(session, todo)
    session.flush()

    default_assignee = project.pm_user_id or project.owner_user_id or current_user.id
    default_due_at = datetime.combine(project.due_at, datetime.min.time()) if project.due_at else None

    created = 0
    for row in expected_rows:
        feature_key = row["feature_key"]
        keep = preserved.get(feature_key, {})
        todo = TodoItem(
            our_entity_id=project.our_entity_id,
            assignee_user_id=keep.get("assignee_user_id") or default_assignee,
            creator_user_id=project.pm_user_id or current_user.id,
            title=row["title"],
            description=row["description"],
            source_type=TodoSourceType.PROJECT_TASK,
            source_id=str(project_id),
            action_type=TodoActionType.DO,
            priority=keep.get("priority") or TodoPriority.P2,
            status=TodoStatus.OPEN,
            due_at=keep.get("due_at") or default_due_at,
            tags=[row["dev_type"]],
            link={
                "project_id": str(project_id),
                "project_name": project.name,
                "dev_type": row["dev_type"],
                "feature_key": feature_key,
                "generated_from_feature_list": True,
                "l1_feature": row["l1_feature"],
                "l2_feature": row["l2_feature"],
                "module": row["module"],
            },
        )
        session.add(todo)
        created += 1

    session.commit()

    return success_response({
        "created": created,
        "deleted": deleted,
        "feature_total": len(expected_rows),
        "source_stage": source_stage.stage_name if source_stage else None,
    })


def sync_project_progress(session, project_id: UUID):
    """Recalculates and saves project progress percentage."""
    from app.models.todo import TodoItem, TodoSourceType
    project = session.get(Project, project_id)
    if not project:
        return
    if project.project_type and project.project_type.lower() == "b2b":
        stages = session.exec(select(ProjectStage).where(ProjectStage.project_id == project_id)).all()
        if not stages:
            project.progress_percentage = 0
        else:
            done = sum(1 for s in stages if s.status == StageStatus.DONE)
            skipped = sum(1 for s in stages if s.status == StageStatus.SKIPPED)
            project.progress_percentage = int(((done + skipped) / len(stages)) * 100)
    else:
        todos = session.exec(
            select(TodoItem)
            .where(TodoItem.source_type == TodoSourceType.PROJECT_TASK)
            .where(TodoItem.source_id == str(project_id))
        ).all()
        if not todos:
            project.progress_percentage = 0
        else:
            done = sum(1 for t in todos if t.status == 'done')
            project.progress_percentage = int((done / len(todos)) * 100)
    session.add(project)
    session.commit()


@router.post("/export_quote_excel")

async def export_quote_excel(
    payload: Dict[str, Any] = Body(...),
    current_user: User = Depends(get_current_user)
):
    """
    Generate a 2-sheet styled Excel: Sheet1=报价单, Sheet2=功能清单
    Payload: { project_name, rows, notes, total_price, total_final, final_confirmed, feature_list }
    """
    project_name = payload.get("project_name", "项目")
    rows = payload.get("rows", [])
    notes = payload.get("notes", [])
    total_price = payload.get("total_price", 0)
    total_final = payload.get("total_final", 0)
    final_confirmed = payload.get("final_confirmed", "")
    feature_list = payload.get("feature_list", [])

    wb = Workbook()

    # ─── Shared Styles ────────────────────────────────────────────────────────
    header_fill = PatternFill(start_color="1F497D", end_color="1F497D", fill_type="solid")
    subtotal_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    total_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
    hours_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
    note_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    white_font = Font(color="FFFFFF", bold=True)
    red_font = Font(color="C00000", bold=True)
    bold_font = Font(bold=True)
    bold_italic = Font(bold=True, italic=True)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    right_align = Alignment(horizontal="right", vertical="center")

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    CHINESE_NUMS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
                    "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十"]

    COL_COUNT = 8  # 序号,服务内容,规格,单位,产品总价,折扣率,折后价,备注

    # ─── SHEET 1: 报价单 ────────────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "报价单"

    # Row 1: Title (styled same as feature list)
    ws1.merge_cells("A1:H1")
    ws1["A1"] = f"{project_name}项目报价"
    ws1["A1"].fill = header_fill
    ws1["A1"].font = white_font
    ws1["A1"].alignment = center_align

    # Row 2: Unit note (right-aligned, plain)
    ws1.merge_cells("A2:G2")
    ws1["H2"] = "单位：人民币（元）"
    ws1["H2"].alignment = right_align
    ws1["H2"].font = Font(size=9, color="666666")

    # Row 3: Headers (dark blue)
    headers = ["序号", "服务内容", "规格", "单位", "产品总价", "折扣率(%)", "折后价", "备注"]
    for ci, h in enumerate(headers, 1):
        cell = ws1.cell(row=3, column=ci, value=h)
        cell.fill = header_fill
        cell.font = white_font
        cell.alignment = center_align
        cell.border = thin_border

    # Column widths
    ws1.column_dimensions["A"].width = 8
    ws1.column_dimensions["B"].width = 30
    ws1.column_dimensions["C"].width = 8
    ws1.column_dimensions["D"].width = 8
    ws1.column_dimensions["E"].width = 14
    ws1.column_dimensions["F"].width = 12
    ws1.column_dimensions["G"].width = 14
    ws1.column_dimensions["H"].width = 30

    current_row = 4
    chinese_counter = 0

    for row in rows:
        if row.get("is_subtotal"):
            ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
            cell = ws1.cell(row=current_row, column=1, value=row.get("subtotal_label", "小计"))
            cell.font = bold_italic
            cell.alignment = center_align
            for ci, val in enumerate([row.get("total_price", ""), "", row.get("final_price", ""), ""], 5):
                c = ws1.cell(row=current_row, column=ci, value=val if val != "" else None)
                c.font = bold_italic
                c.alignment = right_align if ci in (5, 7) else center_align
            for ci in range(1, COL_COUNT + 1):
                ws1.cell(row=current_row, column=ci).fill = subtotal_fill
                ws1.cell(row=current_row, column=ci).border = thin_border
        else:
            chinese_counter += 1
            label = CHINESE_NUMS[chinese_counter - 1] if chinese_counter <= len(CHINESE_NUMS) else str(chinese_counter)
            total_p = row.get("total_price", "")
            final_p = row.get("final_price", "")
            values = [
                label,
                row.get("name", ""),
                row.get("spec", ""),
                row.get("unit", ""),
                float(total_p) if total_p else None,
                row.get("discount", ""),
                float(final_p) if final_p else None,
                row.get("remark", "")
            ]
            for ci, val in enumerate(values, 1):
                cell = ws1.cell(row=current_row, column=ci, value=val)
                cell.border = thin_border
                if ci in (5, 7):
                    cell.alignment = right_align
                elif ci == 2:
                    cell.alignment = left_align
                else:
                    cell.alignment = center_align
        current_row += 1

    # 总计 row
    ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
    ws1.cell(row=current_row, column=1, value="总计").font = red_font
    ws1.cell(row=current_row, column=1).alignment = center_align
    tp_val = float(total_price) if total_price else None
    tf_val = float(total_final) if total_final else None
    ws1.cell(row=current_row, column=5, value=tp_val).font = red_font
    ws1.cell(row=current_row, column=5).alignment = right_align
    ws1.cell(row=current_row, column=7, value=tf_val).font = red_font
    ws1.cell(row=current_row, column=7).alignment = right_align
    for ci in range(1, COL_COUNT + 1):
        ws1.cell(row=current_row, column=ci).fill = total_fill
        ws1.cell(row=current_row, column=ci).border = thin_border
    current_row += 1

    # 最终确认 row
    ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
    ws1.cell(row=current_row, column=1, value="最终确认").font = red_font
    ws1.cell(row=current_row, column=1).alignment = center_align
    try:
        confirmed_val = float(final_confirmed) if final_confirmed != "" else None
    except (ValueError, TypeError):
        confirmed_val = None
    ws1.cell(row=current_row, column=7, value=confirmed_val).font = red_font
    ws1.cell(row=current_row, column=7).alignment = right_align
    for ci in range(1, COL_COUNT + 1):
        ws1.cell(row=current_row, column=ci).fill = total_fill
        ws1.cell(row=current_row, column=ci).border = thin_border
    current_row += 1

    # ─── Work Hours Section (computed from feature_list) ──────────────────────
    # Sum dev hours from feature_list
    dev_backend_total = 0.0
    dev_frontend_total = 0.0
    dev_ui_total = 0.0
    dev_product_total = 0.0
    for f in feature_list:
        for key, acc in [("dev_backend", "dev_backend_total"), ("dev_frontend", "dev_frontend_total"),
                         ("dev_ui", "dev_ui_total"), ("dev_product", "dev_product_total")]:
            try:
                val = float(f.get(key, 0) or 0)
            except (ValueError, TypeError):
                val = 0.0
            if key == "dev_backend": dev_backend_total += val
            elif key == "dev_frontend": dev_frontend_total += val
            elif key == "dev_ui": dev_ui_total += val
            elif key == "dev_product": dev_product_total += val

    total_dev = dev_backend_total + dev_frontend_total + dev_ui_total + dev_product_total
    import math
    test_hrs = math.ceil(total_dev * 0.1)
    mgmt_hrs = math.ceil(total_dev * 0.1)

    def add_hours_row(r, label, backend, frontend, ui, product, fill=hours_fill, font=None):
        ws1.merge_cells(start_row=r, start_column=1, end_row=r, end_column=4)
        cell = ws1.cell(row=r, column=1, value=label)
        cell.alignment = center_align
        cell.border = thin_border
        if font: cell.font = font
        for ci, val in zip([5, 6, 7, 8], [backend, frontend, ui, product]):
            c = ws1.cell(row=r, column=ci, value=val if val else None)
            c.alignment = center_align
            c.border = thin_border
            if font: c.font = font
        for ci in range(1, COL_COUNT + 1):
            ws1.cell(row=r, column=ci).fill = fill
            if not ws1.cell(row=r, column=ci).border.left.style:
                ws1.cell(row=r, column=ci).border = thin_border

    # Header for work hours section
    ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=COL_COUNT)
    ws1.cell(row=current_row, column=1, value="工时明细").fill = header_fill
    ws1.cell(row=current_row, column=1).font = white_font
    ws1.cell(row=current_row, column=1).alignment = center_align
    ws1.cell(row=current_row, column=1).border = thin_border
    current_row += 1

    # Sub-header row
    hours_headers = ["项目", "后端开发", "前端开发", "UI 设计", "产品规划"]
    ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
    ws1.cell(row=current_row, column=1, value="工时类别").fill = header_fill
    ws1.cell(row=current_row, column=1).font = white_font
    ws1.cell(row=current_row, column=1).alignment = center_align
    ws1.cell(row=current_row, column=1).border = thin_border
    for ci, h in zip([5, 6, 7, 8], ["后端开发(天)", "前端开发(天)", "UI设计(天)", "产品规划(天)"]):
        cell = ws1.cell(row=current_row, column=ci, value=h)
        cell.fill = header_fill
        cell.font = white_font
        cell.alignment = center_align
        cell.border = thin_border
    current_row += 1

    # 开发工时
    add_hours_row(current_row, "开发工时",
                  dev_backend_total or None, dev_frontend_total or None,
                  dev_ui_total or None, dev_product_total or None)
    current_row += 1

    # 测试工时 row (merged cols 5-8)
    ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
    ws1.cell(row=current_row, column=1, value="测试工时 (开发合计×10%，向上取整)").alignment = center_align
    ws1.cell(row=current_row, column=1).border = thin_border
    ws1.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=COL_COUNT)
    ws1.cell(row=current_row, column=5, value=test_hrs if test_hrs else None).alignment = center_align
    ws1.cell(row=current_row, column=5).border = thin_border
    for ci in range(1, COL_COUNT + 1):
        ws1.cell(row=current_row, column=ci).fill = hours_fill
        ws1.cell(row=current_row, column=ci).border = thin_border
    current_row += 1

    # 管理工时 row
    ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
    ws1.cell(row=current_row, column=1, value="管理工时 (开发合计×10%，向上取整)").alignment = center_align
    ws1.cell(row=current_row, column=1).border = thin_border
    ws1.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=COL_COUNT)
    ws1.cell(row=current_row, column=5, value=mgmt_hrs if mgmt_hrs else None).alignment = center_align
    ws1.cell(row=current_row, column=5).border = thin_border
    for ci in range(1, COL_COUNT + 1):
        ws1.cell(row=current_row, column=ci).fill = hours_fill
        ws1.cell(row=current_row, column=ci).border = thin_border
    current_row += 1

    # 合计工时 row
    total_hrs = total_dev + test_hrs + mgmt_hrs
    ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
    ws1.cell(row=current_row, column=1, value="合计工时").font = bold_font
    ws1.cell(row=current_row, column=1).alignment = center_align
    ws1.cell(row=current_row, column=1).border = thin_border
    ws1.merge_cells(start_row=current_row, start_column=5, end_row=current_row, end_column=COL_COUNT)
    ws1.cell(row=current_row, column=5, value=total_hrs if total_hrs else None).font = bold_font
    ws1.cell(row=current_row, column=5).alignment = center_align
    ws1.cell(row=current_row, column=5).border = thin_border
    for ci in range(1, COL_COUNT + 1):
        ws1.cell(row=current_row, column=ci).fill = total_fill
        ws1.cell(row=current_row, column=ci).border = thin_border
    current_row += 2

    # 备注 section
    if notes:
        ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=COL_COUNT)
        ws1.cell(row=current_row, column=1, value="备注").font = bold_font
        ws1.cell(row=current_row, column=1).fill = note_fill
        ws1.cell(row=current_row, column=1).alignment = left_align
        ws1.cell(row=current_row, column=1).border = thin_border
        current_row += 1

        notes_text = "\n".join(f"{i+1}、{n}" for i, n in enumerate(notes))
        row_span = max(len(notes), 3)
        ws1.merge_cells(start_row=current_row, start_column=1, end_row=current_row + row_span - 1, end_column=COL_COUNT)
        cell = ws1.cell(row=current_row, column=1, value=notes_text)
        cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        cell.border = thin_border
        ws1.row_dimensions[current_row].height = max(15 * len(notes), 60)

    # ─── SHEET 2: 功能清单 ──────────────────────────────────────────────────
    if feature_list:
        ws2 = wb.create_sheet(title="功能清单")
        headers2 = ["序号", "产品", "板块", "一级功能", "二级功能", "新增功能说明", "后端开发", "前端开发", "UI设计", "产品规划"]
        ws2.merge_cells("A1:J1")
        ws2["A1"] = f"{project_name}功能清单"
        ws2["A1"].fill = header_fill
        ws2["A1"].font = white_font
        ws2["A1"].alignment = center_align

        for ci, h in enumerate(headers2, 1):
            cell = ws2.cell(row=2, column=ci, value=h)
            cell.fill = header_fill
            cell.font = white_font
            cell.alignment = center_align
            cell.border = thin_border

        for col_letter, width in zip("ABCDEFGHIJ", [8, 15, 15, 20, 20, 50, 12, 12, 12, 12]):
            ws2.column_dimensions[col_letter].width = width

        feat_data_end = 2
        for row_i, f in enumerate(feature_list, 3):
            for ci, key in enumerate(["index", "product", "module", "l1_feature", "l2_feature", "description"], 1):
                ws2.cell(row=row_i, column=ci, value=f.get(key, "")).alignment = center_align
                ws2.cell(row=row_i, column=ci).border = thin_border
            for ci, key in enumerate(["dev_backend", "dev_frontend", "dev_ui", "dev_product"], 7):
                val = f.get(key, "")
                try:
                    val = float(val) if val else None
                except (ValueError, TypeError):
                    pass
                ws2.cell(row=row_i, column=ci, value=val).alignment = center_align
                ws2.cell(row=row_i, column=ci).border = thin_border
            feat_data_end = row_i

        # Footer rows matching feature list export
        def apply_feat_footer(r, label, fg, fh=None, fi=None, fj=None):
            ws2.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
            ws2.cell(row=r, column=1, value=label).alignment = center_align
            for ci in range(1, 11):
                ws2.cell(row=r, column=ci).border = thin_border
            ws2.cell(row=r, column=7, value=fg).alignment = center_align
            if fh: ws2.cell(row=r, column=8, value=fh).alignment = center_align
            if fi: ws2.cell(row=r, column=9, value=fi).alignment = center_align
            if fj: ws2.cell(row=r, column=10, value=fj).alignment = center_align

        dev_row = feat_data_end + 1
        apply_feat_footer(dev_row, "系统开发工时",
                          f"=SUM(G3:G{feat_data_end})",
                          f"=SUM(H3:H{feat_data_end})",
                          f"=SUM(I3:I{feat_data_end})",
                          f"=SUM(J3:J{feat_data_end})")
        test_row = dev_row + 1
        ws2.merge_cells(start_row=test_row, start_column=7, end_row=test_row, end_column=10)
        apply_feat_footer(test_row, "系统测试工时", f"=ROUNDUP(SUM(G{dev_row}:J{dev_row})*0.1,0)")
        mgmt_row = dev_row + 2
        ws2.merge_cells(start_row=mgmt_row, start_column=7, end_row=mgmt_row, end_column=10)
        apply_feat_footer(mgmt_row, "项目管理工时", f"=ROUNDUP(SUM(G{dev_row}:J{dev_row})*0.1,0)")
        total_row = dev_row + 3
        ws2.merge_cells(start_row=total_row, start_column=7, end_row=total_row, end_column=10)
        apply_feat_footer(total_row, "总计工时", f"=SUM(G{dev_row}:J{dev_row})+G{test_row}+G{mgmt_row}")

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    filename = f"{project_name}_报价单.xlsx"
    encoded_filename = urllib.parse.quote(filename)
    return StreamingResponse(out, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"})


# ─── Dev Task Generation Endpoints ───────────────────────────────────────────

@router.get("/projects/{project_id}/feature-list", response_model=dict)
async def get_project_feature_list(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """Get the feature list (功能清单) from this project's stages, used to generate dev tasks."""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    feature_list_data, source_stage = _extract_feature_list_for_project(project_id, session)

    # Get project members for assignment
    members_rows = session.exec(
        select(ProjectMember).where(ProjectMember.project_id == project_id)
    ).all()
    members = []
    for m in members_rows:
        u = session.get(User, m.user_id)
        if u:
            members.append({"id": str(u.id), "display_name": u.display_name, "email": u.email})

    # Also include PM
    if project.pm_user_id:
        pm = session.get(User, project.pm_user_id)
        if pm and not any(m["id"] == str(pm.id) for m in members):
            members.insert(0, {"id": str(pm.id), "display_name": pm.display_name, "email": pm.email, "is_pm": True})

    return success_response({
        "feature_list": feature_list_data,
        "source_stage": source_stage.stage_name if source_stage else None,
        "members": members,
        "project": {
            "id": str(project.id),
            "name": project.name,
            "our_entity_id": str(project.our_entity_id),
            "due_at": str(project.due_at) if project.due_at else None,
        }
    })


class DevTaskItem(BaseModel):
    title: str
    description: Optional[str] = None
    assignee_user_id: UUID
    priority: str = "p2"
    due_at: Optional[str] = None
    dev_type: Optional[str] = None  # backend / frontend / ui / product
    feature_key: Optional[str] = None


class GenerateDevTasksRequest(BaseModel):
    tasks: List[DevTaskItem]


@router.post("/projects/{project_id}/generate-dev-tasks", response_model=dict)
async def generate_dev_tasks(
    project_id: UUID,
    payload: GenerateDevTasksRequest,
    session: Session = Depends(get_session),
    current_user: User = Depends(require_permission("project.write"))
):
    """Batch-create TodoItems from feature list rows for the development stage."""
    from app.models.todo import TodoItem, TodoSourceType, TodoActionType, TodoPriority, TodoStatus
    from datetime import datetime as _dt

    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")

    created_ids = []
    for item in payload.tasks:
        # Parse due_at
        due_at = None
        if item.due_at:
            try:
                due_at = _dt.fromisoformat(item.due_at)
            except ValueError:
                pass

        priority_map = {"p0": TodoPriority.P0, "p1": TodoPriority.P1, "p2": TodoPriority.P2, "p3": TodoPriority.P3}
        priority = priority_map.get(item.priority, TodoPriority.P2)

        source_id = str(project_id)

        todo = TodoItem(
            our_entity_id=project.our_entity_id,
            assignee_user_id=item.assignee_user_id,
            # PM is the unified reviewer for development tasks.
            creator_user_id=project.pm_user_id or current_user.id,
            title=item.title,
            description=item.description or "",
            source_type=TodoSourceType.PROJECT_TASK,
            source_id=source_id,
            action_type=TodoActionType.DO,
            priority=priority,
            status=TodoStatus.OPEN,
            due_at=due_at,
            tags=[item.dev_type] if item.dev_type else [],
            link={
                "project_id": str(project_id),
                "project_name": project.name,
                "dev_type": item.dev_type or "",
                "feature_key": item.feature_key or "",
                "generated_from_feature_list": True if item.feature_key else False,
            }
        )
        session.add(todo)
        session.flush()
        created_ids.append(str(todo.id))

    session.commit()

    return success_response({
        "created": len(created_ids),
        "todo_ids": created_ids
    })


# ─── Contract Canvas Endpoints ────────────────────────────────────────────────

@router.get("/projects/{project_id}/contract-context", response_model=dict)
async def get_contract_context(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """Get enriched project context for contract generation (project info, parties, amount)"""
    from app.models.contract import Contract, Counterparty
    
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
    
    pm_user = session.get(User, project.pm_user_id) if project.pm_user_id else None
    # 乙方 = 我方主体 (our entity)
    our_entity = session.get(OurEntity, project.our_entity_id) if project.our_entity_id else None
    # 甲方 defaults to the project's customer
    customer = session.get(Counterparty, project.customer_id) if project.customer_id else None
    contract = session.get(Contract, project.contract_id) if project.contract_id else None
    
    # If a contract is linked, use its parties; otherwise fall back to customer/our_entity
    contract_party_a = None  # 甲方 from contract
    contract_party_c = None  # 丙方 from contract (optional)
    if contract:
        contract_party_a = session.get(Counterparty, contract.party_a_id) if contract.party_a_id else None
        contract_party_c = session.get(Counterparty, contract.party_c_id) if contract.party_c_id else None
    
    # 甲方: prefer contract party_a, else customer
    party_a_data = contract_party_a or customer
    
    def counterparty_info(cp):
        if not cp:
            return None
        return {
            "name": cp.name,
            "identifier": cp.identifier,  # 统一社会信用代码/税号
            "address": cp.address,
            "bank_name": cp.bank_name,
            "bank_account": cp.bank_account,
            "phone": cp.phone,
        }
    
    def our_entity_info(oe):
        if not oe:
            return None
        return {
            "name": oe.name,
            "legal_name": oe.legal_name,
            "uscc": oe.uscc,  # 统一社会信用代码
            "address": oe.address,
        }
    
    context = {
        "project": {
            "id": str(project.id),
            "name": project.name,
            "project_no": project.project_no,
            "project_type": project.project_type,
            "description": project.description,
            "start_at": str(project.start_at) if project.start_at else None,
            "due_at": str(project.due_at) if project.due_at else None,
        },
        "pm": {
            "name": pm_user.display_name,
            "email": pm_user.email,
        } if pm_user else None,
        # 乙方 = 我方主体
        "our_entity": our_entity_info(our_entity),
        # 甲方 = customer or contract party_a
        "party_a": counterparty_info(party_a_data),
        # 乙方 = our_entity (alias for clarity in system prompt)
        "party_b": our_entity_info(our_entity),
        # 丙方 = optional third party from linked contract
        "party_c": counterparty_info(contract_party_c),
        "contract": {
            "contract_no": contract.contract_no,
            "name": contract.name,
            "amount_total": str(contract.amount_total),
            "currency": contract.currency,
            "sign_date": str(contract.sign_date) if contract.sign_date else None,
            "expire_date": str(contract.expire_date) if contract.expire_date else None,
            "summary": contract.summary,
        } if contract else None,
    }
    
    return success_response(context)


@router.post("/export-contract-docx")
async def export_contract_docx(
    payload: Dict[str, Any] = Body(...),
    current_user: User = Depends(get_current_user)
):
    """Export contract markdown as a Word (.docx) document"""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    content_md = payload.get("content_md", "")
    filename = payload.get("filename", "合同") + ".docx"
    
    doc = DocxDocument()
    
    # Set A4 page margins
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
    
    # Style defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    
    def set_paragraph_font(para, size=None, bold=False, alignment=None):
        """Apply font settings to a paragraph."""
        for run in para.runs:
            run.font.name = '宋体'
            if size:
                run.font.size = Pt(size)
            run.font.bold = bold
        if alignment:
            para.alignment = alignment
    
    def parse_and_add(doc, line: str):
        """Parse a single markdown line and add it to the document."""
        line = line.rstrip()
        
        # Heading levels
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # Strip inline markdown from heading text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            docx_level = min(level, 4)  # docx supports Heading 1-9
            para = doc.add_heading(text, level=docx_level)
            set_paragraph_font(para, size=18 - (level * 2), bold=True)
            return
        
        # Horizontal rule
        if re.match(r'^[-*_]{3,}$', line.strip()):
            doc.add_paragraph('─' * 40)
            return
        
        # Unordered list item
        list_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if list_match:
            text = list_match.group(2)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.name = '宋体'
            return
        
        # Ordered list item
        ordered_match = re.match(r'^\s*\d+\.\s+(.*)', line)
        if ordered_match:
            text = ordered_match.group(1)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            para = doc.add_paragraph(style='List Number')
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.name = '宋体'
            return
        
        # Empty line
        if not line.strip():
            doc.add_paragraph()
            return
        
        # Regular paragraph with inline formatting
        para = doc.add_paragraph()
        # Split by bold/italic patterns
        segments = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|~~.*?~~|`.*?`)', line)
        for seg in segments:
            if seg.startswith('***') and seg.endswith('***'):
                run = para.add_run(seg[3:-3])
                run.bold = True
                run.italic = True
            elif seg.startswith('**') and seg.endswith('**'):
                run = para.add_run(seg[2:-2])
                run.bold = True
            elif seg.startswith('*') and seg.endswith('*') and len(seg) > 2:
                run = para.add_run(seg[1:-1])
                run.italic = True
            elif seg.startswith('~~') and seg.endswith('~~'):
                run = para.add_run(seg[2:-2])
                # strikethrough not directly supported, just add as normal
            elif seg.startswith('`') and seg.endswith('`'):
                run = para.add_run(seg[1:-1])
                run.font.name = 'Courier New'
            elif seg:
                run = para.add_run(seg)
            if seg:
                run.font.size = Pt(11)
                if run.font.name != 'Courier New':
                    run.font.name = '宋体'
        para.paragraph_format.line_spacing = Pt(22)
    
    # Process each line
    lines = content_md.split('\n')
    in_table = False
    table_rows = []
    
    for i, line in enumerate(lines):
        # Detect markdown tables
        stripped = line.strip()
        if '|' in stripped:
            if not in_table:
                in_table = True
                table_rows = []
            # Skip separator rows like |---|---|
            if re.match(r'^[\|\s\-:]+$', stripped):
                continue
            cols = [c.strip() for c in stripped.strip('|').split('|')]
            table_rows.append(cols)
            continue
        
        if in_table:
            # Render the collected table
            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                t = doc.add_table(rows=len(table_rows), cols=num_cols)
                t.style = 'Table Grid'
                for ri, row in enumerate(table_rows):
                    for ci, cell_text in enumerate(row):
                        if ci < num_cols:
                            cell = t.cell(ri, ci)
                            cell.text = cell_text
                            if ri == 0:  # header row: bold
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
            in_table = False
            table_rows = []
            # Process current line (not table)
            parse_and_add(doc, line)
        else:
            parse_and_add(doc, line)
    
    # Flush any remaining table
    if in_table and table_rows:
        num_cols = max(len(r) for r in table_rows)
        t = doc.add_table(rows=len(table_rows), cols=num_cols)
        t.style = 'Table Grid'
        for ri, row in enumerate(table_rows):
            for ci, cell_text in enumerate(row):
                if ci < num_cols:
                    cell = t.cell(ri, ci)
                    cell.text = cell_text
                    if ri == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
    
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    encoded_filename = urllib.parse.quote(filename)
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@router.get("/projects/{project_id}/acceptance-report/download")
async def download_acceptance_report(
    project_id: UUID,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    """Download project acceptance report as docx."""
    project = session.get(Project, project_id)
    if not project:
        raise NotFoundException("未找到项目")
        
    feature_list, _ = _extract_feature_list_for_project(project_id, session)
    
    # Get Customer name
    customer_name = ""
    if project.customer_id:
        customer = session.get(Counterparty, project.customer_id)
        if customer:
            customer_name = customer.name
            
    # Get Our Entity name
    our_entity_name = ""
    if project.our_entity_id:
        entity = session.get(OurEntity, project.our_entity_id)
        if entity:
            our_entity_name = entity.name
            
    # Get Members
    members = []
    project_members = session.exec(
        select(ProjectMember).where(ProjectMember.project_id == project_id)
    ).all()
    
    for pm in project_members:
        user = session.get(User, pm.user_id)
        if user:
            members.append({
                "role_in_project": pm.role_in_project or "",
                "user_name": user.display_name or ""
            })
            
    doc_io = generate_acceptance_report(
        project_name=project.name,
        project_no=project.project_no,
        start_date=project.start_at.strftime("%Y-%m-%d") if project.start_at else "",
        due_date=project.due_at.strftime("%Y-%m-%d") if project.due_at else "",
        customer_name=customer_name,
        our_entity_name=our_entity_name,
        members=members,
        feature_list=feature_list
    )
    
    encoded_filename = urllib.parse.quote(f"{project.name}_验收报告.docx")
    return StreamingResponse(
        doc_io,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )
