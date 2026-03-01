import io
from typing import List, Dict, Any, Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _set_cell_background(cell, color_hex: str):
    """Set background color for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def generate_acceptance_report(
    project_name: str,
    project_no: str,
    start_date: str,
    due_date: str,
    customer_name: str,
    our_entity_name: str,
    members: List[Dict[str, str]],
    feature_list: List[Dict[str, Any]]
) -> io.BytesIO:
    """
    Generate a docx file for Project Acceptance Report.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # Title
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{project_name}\n验收报告")
    title_run.font.name = 'SimHei'
    title.style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_heading('一、 项目信息', level=1)
    
    # Project Info
    p = doc.add_paragraph()
    p.add_run(f"项目名称：{project_name}\n").bold = True
    p.add_run(f"项目编号：{project_no}\n")
    p.add_run(f"计划开始日期：{start_date}\n")
    p.add_run(f"计划结束日期：{due_date}\n")
    
    doc.add_heading('二、 业务描述', level=1)
    
    if not feature_list:
        doc.add_paragraph("（无功能清单记录）")
    else:
        for idx, feature in enumerate(feature_list, 1):
            l1 = feature.get("l1_feature", "")
            l2 = feature.get("l2_feature", "")
            desc = feature.get("description") or feature.get("func_desc") or ""
            doc.add_paragraph(f"{idx}. {l1} - {l2}", style='List Number')
            if desc:
                p_desc = doc.add_paragraph(desc)
                p_desc.style.font.size = Pt(10)
    
    doc.add_heading('三、 验收与测试结果', level=1)
    p_verify = doc.add_paragraph()
    p_verify.add_run("经测试和审查，该项目功能均已符合需求定义。")
    
    # Table of checks
    if feature_list:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '一级模块'
        hdr_cells[1].text = '二级功能'
        hdr_cells[2].text = '验收结果'
        
        for cell in hdr_cells:
            _set_cell_background(cell, "E0E0E0")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
        
        for feature in feature_list:
            row_cells = table.add_row().cells
            row_cells[0].text = feature.get("l1_feature", "")
            row_cells[1].text = feature.get("l2_feature", "")
            row_cells[2].text = "☑ 通过  □ 不通过"
    
    doc.add_heading('四、 验收结论', level=1)
    p_conclusion = doc.add_paragraph()
    p_conclusion.add_run(f"经验收测试和文档审查表明，该项目已经达到项目预期目标，验收人员一致认为该项目（ ☑ 通过 □ 不通过 ）验收。\n\n")
    
    p_date = doc.add_paragraph(f"验收日期：      年    月    日")
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    doc.add_heading('五、 签字确认', level=1)
    
    sig_table = doc.add_table(rows=3, cols=5)
    sig_table.style = 'Table Grid'
    
    # Merge cells for header
    sig_table.cell(0, 0).merge(sig_table.cell(1, 0))
    sig_table.cell(0, 1).merge(sig_table.cell(0, 3))
    sig_table.cell(0, 4).merge(sig_table.cell(1, 4))
    
    sig_table.cell(0, 0).text = "签署部门"
    sig_table.cell(0, 1).text = "负责人签字"
    sig_table.cell(0, 4).text = "签署日期"
    sig_table.cell(1, 1).text = "角色"
    sig_table.cell(1, 2).text = "姓名"
    sig_table.cell(1, 3).text = "签名"
    
    for row in range(2):
        for col in range(5):
            cell = sig_table.cell(row, col)
            _set_cell_background(cell, "002060") # Dark blue as in the template
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    
    # Customer Row
    sig_table.cell(2, 0).text = customer_name or "客户方"
    
    # Our Entity Rows
    # Find PM, Tech Lead, Product, etc.
    roles = ["项目总监", "项目经理", "技术经理", "产品负责人"]
    our_rows_start = 3
    for role in roles:
        row = sig_table.add_row()
        row.cells[1].text = role
        # Try to find a member with this role, else leave blank
        name = ""
        for m in members:
            if role in str(m.get("role_in_project", "")):
                name = m.get("user_name", "")
                break
        row.cells[2].text = name
    
    # Merge "Our Entity" column
    sig_table.cell(our_rows_start, 0).merge(sig_table.cell(our_rows_start + len(roles) - 1, 0))
    sig_table.cell(our_rows_start, 0).text = our_entity_name or "我方公司"
    sig_table.cell(our_rows_start, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Merge Signature date column if needed, or leave individual
    sig_table.cell(our_rows_start, 4).merge(sig_table.cell(our_rows_start + len(roles) - 1, 4))
    
    for row in sig_table.rows[2:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    
    return byte_io
